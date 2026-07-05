"""
Фикстуры для тестов.
"""

import os
import sys
import pytest
from unittest.mock import MagicMock, patch

# Добавляем корневую директорию backend в путь
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))


@pytest.fixture
def sample_pdf_path(tmp_path):
    """Создаёт тестовый PDF файл."""
    pdf_path = tmp_path / "test_document.pdf"
    # Минимальный валидный PDF (пустая страница)
    pdf_content = (
        b"%PDF-1.4\n"
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] >>\nendobj\n"
        b"xref\n0 4\n"
        b"0000000000 65535 f \n"
        b"0000000009 00000 n \n"
        b"0000000058 00000 n \n"
        b"0000000115 00000 n \n"
        b"trailer\n<< /Size 4 /Root 1 0 R >>\n"
        b"startxref\n190\n%%EOF"
    )
    pdf_path.write_bytes(pdf_content)
    return str(pdf_path)


@pytest.fixture
def sample_docx_path(tmp_path):
    """Создаёт тестовый DOCX файл."""
    # Для DOCX нужен валидный zip с определённой структурой
    # Используем python-docx для создания
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph("Тестовый текст для поиска.")
        doc.add_paragraph("Второй абзац документа.")
        docx_path = tmp_path / "test_document.docx"
        doc.save(str(docx_path))
        return str(docx_path)
    except ImportError:
        pytest.skip("python-docx не установлен")


@pytest.fixture
def mock_es_client():
    """Мок Elasticsearch клиента."""
    with patch('app.services.elasticsearch_service.Elasticsearch') as MockES:
        mock_instance = MagicMock()
        mock_instance.ping.return_value = True
        mock_instance.indices.exists.return_value = False
        mock_instance.indices.create.return_value = {"acknowledged": True}
        mock_instance.search.return_value = {
            "hits": {
                "total": {"value": 0},
                "hits": []
            }
        }
        MockES.return_value = mock_instance
        yield mock_instance


@pytest.fixture
def mock_db_session():
    """Мок сессии базы данных."""
    session = MagicMock()
    session.commit = MagicMock()
    session.refresh = MagicMock()
    session.close = MagicMock()
    return session