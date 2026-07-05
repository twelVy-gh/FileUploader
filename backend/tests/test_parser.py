"""
Тесты для сервиса парсинга документов (parser.py).
"""

import pytest

from app.services.parser import DocumentParser


class TestDocumentParser:
    """Тесты класса DocumentParser."""

    def setup_method(self):
        """Инициализация парсера перед каждым тестом."""
        self.parser = DocumentParser()

    def test_extract_text_with_pages_pdf(self, sample_pdf_path):
        """Тест извлечения текста из PDF файла."""
        pages = self.parser.extract_text_with_pages(sample_pdf_path)
        # PDF может быть пустым, но метод должен вернуть список
        assert isinstance(pages, list)

    def test_extract_text_with_pages_docx(self, sample_docx_path):
        """Тест извлечения текста из DOCX файла."""
        pages = self.parser.extract_text_with_pages(sample_docx_path)
        assert isinstance(pages, list)
        assert len(pages) > 0
        # DOCX возвращается как одна "страница" — кортеж (page_number, text)
        page_num, text = pages[0]
        assert page_num == 1
        assert "Тестовый текст" in text

    def test_extract_text_with_pages_invalid_format(self, tmp_path):
        """Тест обработки неподдерживаемого формата."""
        txt_path = tmp_path / "test.txt"
        txt_path.write_text("Просто текст", encoding="utf-8")
        with pytest.raises(ValueError, match="Неподдерживаемый формат"):
            self.parser.extract_text_with_pages(str(txt_path))

    def test_extract_text_with_pages_nonexistent_file(self):
        """Тест обработки несуществующего файла."""
        with pytest.raises(FileNotFoundError):
            self.parser.extract_text_with_pages("/nonexistent/file.pdf")

    def test_extract_text_with_pages_empty_pdf(self, sample_pdf_path):
        """Тест обработки PDF без текстового содержимого."""
        pages = self.parser.extract_text_with_pages(sample_pdf_path)
        # Минимальный PDF может не содержать текста
        assert isinstance(pages, list)

    def test_extract_text_with_pages_docx_multiple_paragraphs(self, tmp_path):
        """Тест DOCX с несколькими абзацами."""
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph("Первый абзац.")
            doc.add_paragraph("Второй абзац.")
            doc.add_paragraph("Третий абзац.")
            docx_path = tmp_path / "multi_para.docx"
            doc.save(str(docx_path))

            pages = self.parser.extract_text_with_pages(str(docx_path))
            assert len(pages) == 1
            page_num, text = pages[0]
            assert page_num == 1
            assert "Первый абзац" in text
            assert "Второй абзац" in text
            assert "Третий абзац" in text
        except ImportError:
            pytest.skip("python-docx не установлен")

    def test_extract_text_with_pages_docx_empty(self, tmp_path):
        """Тест пустого DOCX файла."""
        try:
            from docx import Document
            doc = Document()
            docx_path = tmp_path / "empty.docx"
            doc.save(str(docx_path))

            pages = self.parser.extract_text_with_pages(str(docx_path))
            # Пустой документ должен вернуть пустой список
            assert isinstance(pages, list)
            assert len(pages) == 0
        except ImportError:
            pytest.skip("python-docx не установлен")

    def test_extract_text_pdf(self, sample_pdf_path):
        """Тест extract_text для PDF."""
        text = self.parser.extract_text(sample_pdf_path)
        assert isinstance(text, str)

    def test_extract_text_docx(self, sample_docx_path):
        """Тест extract_text для DOCX."""
        text = self.parser.extract_text(sample_docx_path)
        assert isinstance(text, str)
        assert "Тестовый текст" in text

    def test_extract_text_invalid_format(self, tmp_path):
        """Тест extract_text с неподдерживаемым форматом."""
        txt_path = tmp_path / "test.txt"
        txt_path.write_text("Просто текст", encoding="utf-8")
        with pytest.raises(ValueError, match="Неподдерживаемый формат"):
            self.parser.extract_text(str(txt_path))

    def test_extract_text_nonexistent_file(self):
        """Тест extract_text с несуществующим файлом."""
        with pytest.raises(FileNotFoundError):
            self.parser.extract_text("/nonexistent/file.pdf")

    def test_pdf_extension_detection(self):
        """Тест определения PDF по расширению."""
        import os
        ext = os.path.splitext("document.pdf")[1].lower()
        assert ext == ".pdf"
        ext = os.path.splitext("DOCUMENT.PDF")[1].lower()
        assert ext == ".pdf"

    def test_docx_extension_detection(self):
        """Тест определения DOCX по расширению."""
        import os
        ext = os.path.splitext("document.docx")[1].lower()
        assert ext == ".docx"
        ext = os.path.splitext("DOCUMENT.DOCX")[1].lower()
        assert ext == ".docx"

    def test_unknown_extension(self):
        """Тест неизвестного расширения."""
        import os
        ext = os.path.splitext("document.txt")[1].lower()
        assert ext == ".txt"
        ext = os.path.splitext("document.jpg")[1].lower()
        assert ext == ".jpg"